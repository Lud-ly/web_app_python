import streamlit as st
import tempfile
import os
from typing import Callable, Dict, List, Tuple

# Imports for various file formats
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import openpyxl
from fpdf import FPDF
from pdf2docx import Converter
import csv
import json
import xml.etree.ElementTree as ET

# Conditional import for PyYAML
try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False
    st.warning("PyYAML n'est pas installé. La conversion YAML ne sera pas disponible.")

def convert_pdf_to_word(input_file: tempfile._TemporaryFileWrapper) -> str:
    output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
        tmp_pdf.write(input_file.getvalue())
        tmp_pdf.close()
        cv = Converter(tmp_pdf.name)
        cv.convert(output_file)
        cv.close()
    return output_file

def convert_word_to_pdf(input_file: tempfile._TemporaryFileWrapper) -> str:
    output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
    doc = Document(input_file)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)
    pdf.output(output_file)
    return output_file

def convert_excel_to_pdf(input_file: tempfile._TemporaryFileWrapper) -> str:
    output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
    workbook = openpyxl.load_workbook(input_file)
    pdf = FPDF()
    for sheet in workbook.sheetnames:
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        ws = workbook[sheet]
        pdf.cell(200, 10, txt=f"Sheet: {sheet}", ln=True)
        for row in ws.iter_rows(values_only=True):
            pdf.cell(200, 10, txt=" | ".join(str(cell) for cell in row), ln=True)
    pdf.output(output_file)
    return output_file

def convert_to_text(input_file: tempfile._TemporaryFileWrapper, file_extension: str) -> str:
    if file_extension == '.pdf':
        reader = PdfReader(input_file)
        return "\n".join(page.extract_text() for page in reader.pages)
    elif file_extension == '.docx':
        doc = Document(input_file)
        return "\n".join(para.text for para in doc.paragraphs)
    elif file_extension == '.xlsx':
        workbook = openpyxl.load_workbook(input_file)
        text = []
        for sheet in workbook.sheetnames:
            ws = workbook[sheet]
            text.append(f"Sheet: {sheet}")
            for row in ws.iter_rows(values_only=True):
                text.append(" | ".join(str(cell) for cell in row))
        return "\n".join(text)
    elif file_extension in ['.csv', '.json', '.xml', '.yaml']:
        return input_file.getvalue().decode('utf-8')
    else:
        return "Unsupported file format for text conversion."

def convert_to_csv(input_file: tempfile._TemporaryFileWrapper, file_extension: str) -> str:
    output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.csv').name
    if file_extension == '.xlsx':
        workbook = openpyxl.load_workbook(input_file)
        sheet = workbook.active
        with open(output_file, 'w', newline='') as f:
            writer = csv.writer(f)
            for row in sheet.iter_rows(values_only=True):
                writer.writerow(row)
    elif file_extension == '.json':
        data = json.load(input_file)
        with open(output_file, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(data[0].keys())
            for row in data:
                writer.writerow(row.values())
    elif file_extension == '.xml':
        tree = ET.parse(input_file)
        root = tree.getroot()
        with open(output_file, 'w', newline='') as f:
            writer = csv.writer(f)
            headers = [elem.tag for elem in root[0]]
            writer.writerow(headers)
            for child in root:
                writer.writerow([child.find(elem).text for elem in headers])
    return output_file

def convert_to_json(input_file: tempfile._TemporaryFileWrapper, file_extension: str) -> str:
    output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.json').name
    if file_extension == '.csv':
        with open(input_file.name, newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            data = list(reader)
        with open(output_file, 'w') as jsonfile:
            json.dump(data, jsonfile, indent=2)
    elif file_extension == '.xml':
        tree = ET.parse(input_file)
        root = tree.getroot()
        data = []
        for child in root:
            item = {elem.tag: elem.text for elem in child}
            data.append(item)
        with open(output_file, 'w') as jsonfile:
            json.dump(data, jsonfile, indent=2)
    elif file_extension == '.yaml' and YAML_AVAILABLE:
        data = yaml.safe_load(input_file)
        with open(output_file, 'w') as jsonfile:
            json.dump(data, jsonfile, indent=2)
    return output_file

# Define conversion functions
CONVERSION_FUNCTIONS: Dict[Tuple[str, str], Callable] = {
    ('.pdf', '.docx'): convert_pdf_to_word,
    ('.docx', '.pdf'): convert_word_to_pdf,
    ('.xlsx', '.pdf'): convert_excel_to_pdf,
    ('.pdf', '.txt'): lambda f: convert_to_text(f, '.pdf'),
    ('.docx', '.txt'): lambda f: convert_to_text(f, '.docx'),
    ('.xlsx', '.txt'): lambda f: convert_to_text(f, '.xlsx'),
    ('.csv', '.txt'): lambda f: convert_to_text(f, '.csv'),
    ('.json', '.txt'): lambda f: convert_to_text(f, '.json'),
    ('.xml', '.txt'): lambda f: convert_to_text(f, '.xml'),
    ('.xlsx', '.csv'): lambda f: convert_to_csv(f, '.xlsx'),
    ('.json', '.csv'): lambda f: convert_to_csv(f, '.json'),
    ('.xml', '.csv'): lambda f: convert_to_csv(f, '.xml'),
    ('.csv', '.json'): lambda f: convert_to_json(f, '.csv'),
    ('.xml', '.json'): lambda f: convert_to_json(f, '.xml'),
}

# Add YAML conversions only if PyYAML is available
if YAML_AVAILABLE:
    CONVERSION_FUNCTIONS[('.yaml', '.txt')] = lambda f: convert_to_text(f, '.yaml')
    CONVERSION_FUNCTIONS[('.yaml', '.json')] = lambda f: convert_to_json(f, '.yaml')

# Supported file formats
SUPPORTED_FORMATS: List[str] = ['.pdf', '.docx', '.xlsx', '.txt', '.csv', '.json', '.xml']
if YAML_AVAILABLE:
    SUPPORTED_FORMATS.append('.yaml')

def main():
    st.title("Mon Convertisseur de Fichiers")
    st.write("Ce convertisseur de fichiers polyvalent permet de transformer facilement des documents entre différents formats courants, incluant PDF, Word, Excel, texte brut, CSV, JSON, XML et YAML (si disponible), offrant ainsi une solution pratique pour la conversion de fichiers dans un environnement web interactif.")
    st.image("./images/convert.jpg", width=200)

    # Afficher les formats pris en charge
    st.write("Formats pris en charge :", ", ".join(SUPPORTED_FORMATS))

    uploaded_file = st.file_uploader("Choisissez un fichier...", type=[fmt[1:] for fmt in SUPPORTED_FORMATS])

    if uploaded_file is not None:
        input_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        # Filter available conversion formats based on input file type
        available_formats = [fmt for fmt in SUPPORTED_FORMATS if (input_extension, fmt) in CONVERSION_FUNCTIONS]
        
        conversion_format = st.selectbox(
            "Convertir en :",
            options=available_formats
        )

        if st.button("Convertir"):
            conversion_key = (input_extension, conversion_format)
            if conversion_key in CONVERSION_FUNCTIONS:
                st.write(f"Conversion de {input_extension} à {conversion_format}...")
                try:
                    output_file = CONVERSION_FUNCTIONS[conversion_key](uploaded_file)
                    st.success(f"Fichier converti avec succès !")
                    
                    if conversion_format == '.txt':
                        with open(output_file, 'r') as f:
                            text_content = f.read()
                        st.text_area("Aperçu du contenu converti :", value=text_content[:1000] + "...", height=300)
                    
                    # Bouton de téléchargement pour tous les formats
                    with open(output_file, 'rb') as f:
                        st.download_button(
                            label="Télécharger le fichier converti",
                            data=f,
                            file_name=f"converted{conversion_format}",
                            mime=f"application/{conversion_format[1:]}"
                        )
                    
                except Exception as e:
                    st.error(f"Erreur lors de la conversion : {str(e)}")
            else:
                st.error("Cette conversion n'est pas prise en charge.")

if __name__ == "__main__":
    main()